import docx
import csv

for i in range(22):
    if i <= 8:
        st = '0'+str(i)
        end = '0'+str(i+1)
    elif i == 9:
        st = '0' + str(i)
        end = str(i+1)
    else:
        st = str(i)
        end = str(i+1)
    team_per_game = './nba_data/{}-{} team per game.docx'.format(st,end)
    player_per_game = './nba_data/{}-{} player per game.docx'.format(st,end)
    #team_advanced = './nba_data/{}-{} team advanced stats.docx'.format(st,end)
    opp_per_game = './nba_data/{}-{} opponent per game.docx'.format(st,end)

    #ignore advanced stats for now
    #loop thru each doc and create a new csv file
        #first line in each doc is column names
        #loop thru each line and write to csv file
    


    doc = docx.Document(team_per_game)
    tpgText = []
    for docpara in doc.paragraphs:
        row = docpara.text.split(',')
        if len(row)>1:
            row[1] = row[1].replace('*','')
        tpgText.append(row)
    tpgCols = tpgText[0]
    tpgText = tpgText[1:]
    with open('./nba_data/{}-{}_team_per_game.csv'.format(st,end), 'w') as file:
        writer = csv.writer(file)
        writer.writerow(tpgCols)
        writer.writerows(tpgText)
    
    doc = docx.Document(player_per_game)
    ppgText = []
    for docpara in doc.paragraphs:
        row = docpara.text.split(',')
        ppgText.append(row)
    ppgCols = ppgText[0]
    ppgText = ppgText[1:]

    with open('./nba_data/{}-{}_player_per_game.csv'.format(st,end), 'w') as file:
        writer = csv.writer(file)
        writer.writerow(ppgCols)
        writer.writerows(ppgText)

    doc = docx.Document(opp_per_game)
    opgText = []
    for docpara in doc.paragraphs:
        row = docpara.text.split(',')
        if len(row)>1:
            row[1] = row[1].replace('*','')
        opgText.append(row)
    opgCols = opgText[0]
    opgText = opgText[1:]

    with open('./nba_data/{}-{}_opponent_per_game.csv'.format(st,end), 'w') as file:
        writer = csv.writer(file)
        writer.writerow(opgCols)
        writer.writerows(opgText)


   
